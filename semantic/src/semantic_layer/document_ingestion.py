"""Document ingestion pipeline: parse, chunk, embed, store."""

from __future__ import annotations

import asyncio
import hashlib
import uuid
from pathlib import Path
from typing import TYPE_CHECKING

import structlog
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import text

from seld_common.db import DatabaseManager
from semantic_layer.embeddings import EmbeddingService

if TYPE_CHECKING:
    from semantic_layer.config import Settings

logger = structlog.get_logger(__name__)

class DocumentIngestion:
    """Ingests documents into the vector store for RAG retrieval.

    Two-table schema:
    - ``documents``: one row per ingested file (metadata + status).
    - ``document_chunks``: one row per chunk (content + embedding + FK).
    """

    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._db = DatabaseManager(settings.database_url)
        self._embedder = EmbeddingService(settings)
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            add_start_index=True, # add the index of the chunk in the document
        )

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    async def ingest_text(
        self,
        content: str,
        source: str,
        asset_tags: list[str] | None = None,
    ) -> int:
        """Ingest raw text: split → embed → store.

        Returns:
            Number of chunks stored.
        """
        chunks = self._splitter.split_text(content)
        if not chunks:
            return 0

        pages: list[int | None] = [None] * len(chunks)
        embeddings = await self._embed_batched(chunks)

        await self._store(
            chunks=chunks,
            pages=pages,
            embeddings=embeddings,
            source=source,
            file_type="txt",
            file_size_bytes=len(content.encode()),
            asset_tags=asset_tags,
        )
        return len(chunks)

    async def ingest_pdf(
        self,
        file_path: str,
        asset_tags: list[str] | None = None,
    ) -> int:
        """Ingest a PDF: extract full text, chunk globally, remap page numbers.

        Uses PyMuPDF (fitz) for robust extraction of multi-column layouts,
        tables, and watermarked industrial manuals.  Chunking runs on the
        *entire document text* so cross-page paragraphs are never broken.
        Each chunk is then mapped back to its originating page via
        character-offset tracking.

        Returns:
            Number of chunks stored.
        """
        import fitz  # PyMuPDF

        path = Path(file_path)
        doc = fitz.open(str(path))

        # ── 1. Extract per-page text + track character offsets ─────────
        page_texts: list[str] = []
        page_offsets: list[tuple[int, int, int]] = []  # (start, end, page_num)
        offset = 0

        for page_idx in range(len(doc)):
            page_text = doc[page_idx].get_text("text")
            if not page_text or not page_text.strip():
                continue
            page_num = page_idx + 1
            page_texts.append(page_text)
            page_offsets.append((offset, offset + len(page_text), page_num))
            offset += len(page_text) + 1  # +1 for "\n" join separator

        doc.close()

        if not page_texts:
            logger.warning("pdf_no_text", file=file_path)
            return 0

        full_text = "\n".join(page_texts)
        if page_texts:
            assert offset - 1 == len(full_text), f"Page offset mismatch: {offset - 1} != {len(full_text)}"

        # ── 2. Chunk the full document ─────────────────────────────────
        from langchain_core.documents import Document
        lc_doc = Document(page_content=full_text)
        lc_chunks = self._splitter.split_documents([lc_doc])
        
        if not lc_chunks:
            return 0

        chunks_data = [
            {"content": c.page_content, "start_index": c.metadata.get("start_index", 0)} 
            for c in lc_chunks
        ]
        text_chunks = [c["content"] for c in chunks_data]

        # ── 3. Map each chunk → page number ───────────────────────────
        pages = self._map_chunks_to_pages(chunks_data, page_offsets)

        # ── 4. Embed + store ──────────────────────────────────────────
        embeddings = await self._embed_batched(text_chunks)

        await self._store(
            chunks=text_chunks,
            pages=pages,
            embeddings=embeddings,
            source=path.name,
            file_type="pdf",
            file_size_bytes=path.stat().st_size,
            asset_tags=asset_tags,
        )

        logger.info("pdf_ingested", file=file_path, total_chunks=len(text_chunks))
        return len(text_chunks)

    async def ingest_docx(
        self,
        file_path: str,
        asset_tags: list[str] | None = None,
    ) -> int:
        """Ingest a DOCX file (no meaningful page numbers)."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not full_text.strip():
            return 0
        return await self.ingest_text(full_text, Path(file_path).name, asset_tags)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    async def _embed_batched(self, chunks: list[str]) -> list[list[float]]:
        """Embed chunks in batches, offloading blocking calls to a thread pool.

        FIX #4: run_in_executor prevents blocking the async event loop.
        FIX #5: batch sizing respects API payload limits.
        """
        loop = asyncio.get_running_loop()
        all_embeddings: list[list[float]] = []
        batch_size = self._settings.embed_batch_size

        for start in range(0, len(chunks), batch_size):
            batch = chunks[start : start + batch_size]

            raw = await loop.run_in_executor(
                None, self._embedder.embed, batch,
            )

            # Normalise numpy arrays → plain lists for pgvector
            for emb in raw:
                all_embeddings.append(
                    emb.tolist() if hasattr(emb, "tolist") else list(emb)
                )

        return all_embeddings

    async def _store(
        self,
        *,
        chunks: list[str],
        pages: list[int | None],
        embeddings: list[list[float]],
        source: str,
        file_type: str,
        file_size_bytes: int,
        asset_tags: list[str] | None,
    ) -> None:
        #Persist document + all chunks in a single transaction.
        document_id = str(uuid.uuid4())
        tags = asset_tags or []

        async with self._db.session() as session:
            # ── Parent document row ────────────────────────────────────
            await session.execute(
                text("""
                    INSERT INTO documents
                        (id, file_name, file_type, file_size_bytes,
                         asset_tags, num_chunks, status)
                    VALUES
                        (:id, :file_name, :file_type, :file_size_bytes,
                         :asset_tags, 0, 'processing')
                """),
                {
                    "id": document_id,
                    "file_name": source,
                    "file_type": file_type,
                    "file_size_bytes": file_size_bytes,
                    "asset_tags": tags,
                },
            )

            # ── Batch insert all chunks ────────────────────────────────
            chunk_params = [
                {
                    "document_id": document_id,
                    "chunk_index": i,
                    "content": chunk,
                    "source": source,
                    "page": page,
                    "content_hash": hashlib.sha256(chunk.encode()).hexdigest(),
                    "asset_tags": tags,
                    "embedding": embedding,
                }
                for i, (chunk, page, embedding) in enumerate(
                    zip(chunks, pages, embeddings)
                )
            ]

            if chunk_params:
                await session.execute(
                    text("""
                        INSERT INTO document_chunks
                            (document_id, chunk_index, content, source,
                             page, content_hash, asset_tags, embedding)
                        VALUES
                            (:document_id, :chunk_index, :content, :source,
                             :page, :content_hash, :asset_tags, :embedding::vector)
                        ON CONFLICT (content_hash) DO NOTHING
                    """),
                    chunk_params,
                )
            await session.execute(
                text("""
                    UPDATE documents
                    SET 
                        num_chunks = (SELECT COUNT(*) FROM document_chunks WHERE document_id = :id),
                        status = 'ready'
                    WHERE id = :id
                """),
                {"id": document_id},
            )

        logger.info(
            "document_ingested",
            source=source,
            file_type=file_type,
            chunks=len(chunks),
        )

    @staticmethod
    def _map_chunks_to_pages(
        chunks_with_offsets: list[dict],
        page_offsets: list[tuple[int, int, int]],
    ) -> list[int | None]:
        """Map each chunk back to its PDF page via character-offset lookup."""
        pages: list[int | None] = []
        for chunk in chunks_with_offsets:
            start = chunk["start_index"]
            page: int | None = None
            for page_start, page_end, pg_num in page_offsets:
                if page_start <= start < page_end:
                    page = pg_num
                    break
            pages.append(page)
        return pages

    async def delete_document(self, document_id: str) -> bool:
        """Deletes a document and all its chunks (via CASCADE) from the database."""
        async with self._db.session() as session:
            result = await session.execute(
                text("DELETE FROM documents WHERE id = :id RETURNING id"),
                {"id": document_id},
            )
            return result.rowcount > 0